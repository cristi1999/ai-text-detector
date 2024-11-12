import time
import streamlit as st
from docx import Document
from pathlib import Path
import pickle
from transformers import pipeline, AutoModelForSequenceClassification, AutoTokenizer
import sys
import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

sys.path.append("..")
import text_utils

MIN_PARAGRAPH_SIZE = 250


LR_PATH = Path("models/lr_classifier.pkl")
XGB_PATH = Path("models/xgb_classifier.pkl")
RF_PATH = Path("models/rf_classifier.pkl")
BERT_PATH = Path("models/fine-tuned-distilbert-base-uncased")
DISTILBERT_PATH = Path("models/fine-tuned-distilbert-base-uncased")
ROBERTA_PATH = Path("models/fine-tuned-roberta-base")
DISTILROBERTA_PATH = Path("models/fine-tuned-distilroberta-base")


def load_llm_classifier(model_path):
    model = AutoModelForSequenceClassification.from_pretrained(model_path)
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    classifier = pipeline("text-classification", model=model, tokenizer=tokenizer)
    return classifier


def is_paragraph_empty(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return False
    return True


def add_shading(paragraph, color="EE0D49"):
    p = paragraph._element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)

    p.get_or_add_pPr().append(shd)


def convert_to_docx(text):
    paragraphs = text.split("\n")
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    return doc


def preprocess_text(text):
    text = " ".join(text.splitlines())
    cleaned_text = " ".join(text_utils.clean_text(text).splitlines())
    word_count = text_utils.count_words(text)
    char_count = text_utils.count_chars(text)
    punctuation_count = text_utils.count_punctuation(text)
    sentence_count = text_utils.count_sentences(text)
    avg_word_length = text_utils.avg_word_length(text)
    unique_word_count = text_utils.count_unique_words(text)
    stopword_count = text_utils.count_stopwords(text)
    avg_sentence_length = text_utils.avg_sentence_length(text)
    hapax_legomena = text_utils.hapax_legomena(text)
    type_token_ratio = text_utils.type_token_ratio(text)
    return pd.DataFrame(
        {
            "text": cleaned_text,
            "word_count": word_count,
            "char_count": char_count,
            "punctuation_count": punctuation_count,
            "sentence_count": sentence_count,
            "avg_word_length": avg_word_length,
            "unique_word_count": unique_word_count,
            "stopword_count": stopword_count,
            "avg_sentence_length": avg_sentence_length,
            "hapax_legomena": hapax_legomena,
            "type_token_ratio": type_token_ratio,
        },
        index=[0],
    )


def load_ml_classifier(model_path):
    with open(model_path, "rb") as file:
        model = pickle.load(file)
    return model


def read_txt(file):
    return file.read().decode("utf-8")


def read_docx(file):
    doc = Document(file)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


bert_classifier = load_llm_classifier(BERT_PATH)
distilbert_classifier = load_llm_classifier(DISTILBERT_PATH)
roberta_classifier = load_llm_classifier(ROBERTA_PATH)
distilroberta_classifier = load_llm_classifier(DISTILROBERTA_PATH)

lr_classifier = load_ml_classifier(LR_PATH)
xgb_classifier = load_ml_classifier(XGB_PATH)
rf_classifier = load_ml_classifier(RF_PATH)

st.set_page_config(page_title="AI Text Detector", page_icon=":detective:")


def detect_text(model_name, text):
    match model_name:
        case "bert":
            classifier = bert_classifier
        case "distilbert":
            classifier = distilbert_classifier
        case "roberta":
            classifier = roberta_classifier
        case "distilroberta":
            classifier = distilroberta_classifier
        case "logistic_regression":
            classifier = lr_classifier
        case "xgboost":
            classifier = xgb_classifier
        case "random_forest":
            classifier = rf_classifier
    if model_name in ["logistic_regression", "xgboost", "random_forest"]:
        input_df = preprocess_text(text)
        prediction = classifier.predict_proba(input_df)
        probability = prediction[0][1]
    else:
        prediction = classifier(text)
        if prediction[0]["label"] == "LABEL_1":
            probability = prediction[0]["score"]
        else:
            probability = 1 - prediction[0]["score"]
    return probability


st.markdown(
    """
## 🎭 TruAIText
"""
)
st.caption("Detect AI generated text using different AI models")

st.markdown("------")

# Text input
with st.expander("Enter text 📝", expanded=False):
    text_input = st.text_area(
        "Input",
        height=300,
        placeholder="Type or paste your text here...",
        label_visibility="collapsed",
    )
    if len(text_input) < 50:
        allow_analyze = False
        st.error("Please enter at least 50 characters!")
    else:
        allow_analyze = True

with st.expander("Upload a file 📁", expanded=False):
    uploaded_file = st.file_uploader("Choose a file", type=["txt", "docx"])
    if uploaded_file:
        file_type = uploaded_file.type
        if file_type == "text/plain":
            text_input = read_txt(uploaded_file)
            st.text_area("File Content", value=text_input, height=300)
        elif (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            text_input = read_docx(uploaded_file)
            st.text_area("File Content", value=text_input, height=300)
        else:
            st.error("Unsupported file type")
        allow_analyze = True

# Model selection
model_options = [
    ("Logistic Regression", "logistic_regression"),
    ("XGBoost", "xgboost"),
    ("Random Forest", "random_forest"),
    ("DistilBERT", "distilbert"),
    ("BERT", "bert"),
    ("DistilRoBERTa", "distilroberta"),
    ("RoBERTa", "roberta"),
]

st.markdown(
    """
**Select the model to use for text detection**
"""
)


def on_change():
    st.session_state["analyze_button"] = False


col_1, col_2 = st.columns([2, 1])
with col_1:
    selected_model = st.selectbox(
        label="Choose a model",
        options=model_options,
        index=0,
        format_func=lambda x: x[0],
        label_visibility="collapsed",
        on_change=on_change,
    )[1]

with col_2:
    analyze_button = st.button(
        "Detect", use_container_width=True, disabled=allow_analyze is False
    )
    if analyze_button:
        st.session_state["analyze_button"] = True


if "analyze_button" in st.session_state and st.session_state["analyze_button"] == True:
    if text_input:
        with st.spinner("Detecting..."):
            overall_probability = detect_text(selected_model, text_input)

            document = convert_to_docx(text_input)
            report = Document()
            report.add_heading("AI Text Detection Paragraph Level Report", level=1)

            probabilities = []
            num_paragraphs = 0
            for paragraph in document.paragraphs:
                text_input = paragraph.text
                if not is_paragraph_empty(paragraph):
                    if len(text_input) >= MIN_PARAGRAPH_SIZE:
                        num_paragraphs += 1
                        probability = detect_text(selected_model, text_input)
                        probabilities.append(probability)
                        report.add_heading("This paragraph has AI-generated percentage of " + str(int(probability * 100)) + "%", level=3)
                        p = report.add_paragraph(text_input)
                        if probability > 0.5:
                            add_shading(p)
                    else:
                        report.add_heading("This paragraph is too short to analyze", level=3)
                        report.add_paragraph(text_input)
            if num_paragraphs > 0:
                doc_probability = sum(probabilities) / num_paragraphs
                report.add_heading(
                    f"Overall Paragraph Level AI Generated Text Percentage: {int(doc_probability * 100)}%",
                    level=2,
                )
            else:
                report.add_heading("No paragraphs detected", level=2)

        st.success("Detection completed!")

        st.write("#### Results")
        percentage = int(overall_probability * 100)
        st.write(f"AI generated text percentage: {percentage}%")
        st.progress(int(percentage))
        st.write("")

        bytes_io = BytesIO()
        report.save(bytes_io)

        if st.download_button(
            label="Download report",
            data=bytes_io.getvalue(),
            file_name="report.docx",
            mime="docx",
        ):
            st.toast("Downloaded report successfully!")
    else:
        st.warning("Please enter some text to analyze.")

# Footer
st.markdown("---")
st.write("Copyright © 2024 Cristian Grecu. All rights reserved.")
